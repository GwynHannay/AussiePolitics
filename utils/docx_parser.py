import docx


def open_document(filename: str):
    """Opens and returns a docx object (Microsoft Word)

    Args:
        filename (str): Full filename and path of docx file.

    Returns:
        |Document| object: docx object to be parsed.
    """
    return docx.Document(filename)


def get_document_text(document) -> list:
    """Grabs and returns every paragraph in a docx document in a list.

    Args:
        document (|Document| object): docx object.

    Returns:
        list: Text of each paragraph from the docx object as a list item.
    """
    document_text = []
    for paragraph in document.paragraphs:
        document_text.append(paragraph.text)

    return document_text


def get_header_text(document) -> list:
    """Grabs and returns the text of each section header from a docx object.

    Args:
        document (|Document| object): docx object.

    Returns:
        list: Text from the header of each section in the docx object as a list item.
    """
    header_text = []
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            header_text.append(paragraph.text)

    return header_text


def get_footer_text(document) -> list:
    """Grabs and returns the text of each section footer from a docx object.

    Args:
        document (|Document| object): docx object.

    Returns:
        list: Text from the footer of each section in the docx object as a list item.
    """
    footer_text = []
    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            footer_text.append(paragraph.text)

    return footer_text


def get_table_text(document) -> list:
    """Grabs and returns the text from each cell of each table from a docx object.

    Args:
        document (|Document| object): docx object.

    Returns:
        list: Text from each cell of each row of each table in the docx object as a list inside a list inside a list.
    """
    table_text = []
    for tab in document.tables:
        row_text = []
        for row in tab.rows:
            cell_text = []
            for cell in row.cells:
                cell_text.append(cell.text)
            row_text.append(cell_text)
        table_text.append(row_text)

    return table_text


def get_inline_shapes(document) -> list:
    """Grabs and returns the type of each inline shape in a docx object.

    Args:
        document (|Document| object): docx object.

    Returns:
        list: Each type of inline shape found in a docx object as a list item.
    """    
    shapes = []
    for shape in document.inline_shapes:
        shapes.append(shape.type)

    return shapes
